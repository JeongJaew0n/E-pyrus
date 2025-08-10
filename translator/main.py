import openai
from docx import Document
import time
from pathlib import Path

# ✅ OpenAI API 키 설정
OPENAI_API_KEY = ""  # <-- 여기에 본인의 API 키 입력

# ✅ 현재 스크립트 기준 상대 경로로 입력/출력 경로 지정
BASE_DIR = Path(__file__).resolve().parent
INPUT_PATH = BASE_DIR / "resources" / "input.docx"
OUTPUT_PATH = BASE_DIR / "translated_output.docx"

# ✅ 번역할 텍스트 분할 함수 (단락 단위)
def load_docx_paragraphs(file_path):
    if not file_path.exists():
        raise FileNotFoundError(f"입력 파일을 찾을 수 없습니다: {file_path}")
    doc = Document(file_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# ✅ GPT를 이용한 자연스러운 번역 함수 (openai>=1.0.0 기준)
def translate_paragraph(paragraph):
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful and natural English-to-Korean translator."},
                {"role": "user", "content": f"Translate the following English paragraph into natural and fluent Korean, preserving the tone and nuances.\n\n{paragraph}"}
            ],
            temperature=0.4,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("Error:", e)
        return "[번역 오류 발생]"

# ✅ 결과 저장 함수
def save_to_docx(paragraphs, output_path):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(output_path)

# ✅ 전체 파이프라인 실행 함수
def translate_docx_file():
    paragraphs = load_docx_paragraphs(INPUT_PATH)
    translated = []

    for i, para in enumerate(paragraphs):
        print(f"[{i+1}/{len(paragraphs)}] 번역 중...")
        kr = translate_paragraph(para)
        translated.append(kr)
        time.sleep(1.2)

    save_to_docx(translated, OUTPUT_PATH)
    print(f"\n✅ 번역 완료: {OUTPUT_PATH}")

# 🔽 실행 (직접 실행 시)
if __name__ == "__main__":
    translate_docx_file()